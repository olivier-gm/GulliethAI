import subprocess
import os
import string
import random
import shutil
import logging
import threading
import unicodedata
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class Document_process:

    # LibreOffice es de instancia única: las conversiones se serializan.
    _LO_LOCK = threading.Lock()

    # ── LibreOffice detection ──────────────────────────────────────────

    @staticmethod
    def _find_libreoffice():
        """Busca la ruta de LibreOffice en ubicaciones comunes."""
        possible_paths = [
            'C:/Program Files/LibreOffice/program/soffice.exe',
            'C:/Program Files (x86)/LibreOffice/program/soffice.exe',
            '/usr/bin/soffice',
            '/usr/local/bin/soffice',
        ]
        for path in possible_paths:
            if os.path.isfile(path):
                return path
        # Intentar con PATH del sistema
        found = shutil.which('soffice')
        if found:
            return found
        raise FileNotFoundError(
            "LibreOffice no encontrado. Instálalo o agrega 'soffice' al PATH."
        )

    @staticmethod
    def _find_libreoffice_python():
        """Busca el intérprete Python que trae LibreOffice (el que tiene `uno`).

        Es un intérprete distinto al del proyecto: el módulo `uno` sólo existe
        dentro de la instalación de LibreOffice.
        """
        try:
            libre_path = Document_process._find_libreoffice()
        except FileNotFoundError:
            return None

        program_dir = os.path.dirname(libre_path)
        candidates = [
            os.path.join(program_dir, 'python.exe'),   # Windows
            os.path.join(program_dir, 'python'),       # algunos builds
            '/usr/bin/python3',                        # Linux con python3-uno
        ]
        for candidate in candidates:
            if os.path.isfile(candidate):
                return candidate
        return None

    @staticmethod
    def _finalize_with_uno(input_file, pdf_output):
        """Actualiza el índice y exporta el PDF vía UNO.

        Devuelve True si LibreOffice pudo poblar el índice y generar el PDF.
        """
        lo_python = Document_process._find_libreoffice_python()
        if not lo_python:
            logger.info('No se encontró el Python de LibreOffice; se usa la conversión simple.')
            return False

        script = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'lo_finalize.py')
        if not os.path.isfile(script):
            logger.warning('No se encontró lo_finalize.py; se usa la conversión simple.')
            return False

        soffice_path = Document_process._find_libreoffice()
        command = [
            lo_python, script,
            os.path.abspath(input_file), os.path.abspath(pdf_output), soffice_path,
        ]

        # Una sola conversión a la vez: LibreOffice es de instancia única y dos
        # procesos simultáneos se bloquean entre sí.
        # SAL_DISABLE_UPDATECHECK=1: suprime el diálogo de actualización de
        # LibreOffice. Sin esto, si hay conexión lenta o actualización pendiente,
        # el diálogo bloquea el proceso headless y la conversión falla.
        env = os.environ.copy()
        env['SAL_DISABLE_UPDATECHECK'] = '1'

        with Document_process._LO_LOCK:
            try:
                result = subprocess.run(
                    command, timeout=150, capture_output=True, text=True, env=env,
                )
            except subprocess.TimeoutExpired:
                logger.error('Timeout actualizando el índice con LibreOffice: %s', input_file)
                return False
            except Exception as e:
                logger.error('Error lanzando LibreOffice/UNO: %s', e)
                return False

        if result.returncode == 0 and os.path.isfile(pdf_output):
            logger.info('Índice actualizado y PDF generado: %s', pdf_output)
            return True

        logger.warning(
            'La finalización con UNO falló (código %s): %s',
            result.returncode, (result.stderr or '').strip()[:500],
        )
        return False

    @staticmethod
    def convert(input_file, output_folder):
        """Genera el PDF y deja el índice poblado en ambos formatos.

        Primero intenta la vía UNO, que además de convertir actualiza la tabla
        de contenido (LibreOffice no lo hace al cargar, así que el índice
        saldría vacío tanto en el PDF como en el Word). Si esa vía no está
        disponible, cae a la conversión simple de siempre: el PDF se genera
        igual, sólo que con el índice sin poblar.
        """
        pdf_output = os.path.join(
            output_folder,
            os.path.splitext(os.path.basename(input_file))[0] + '.pdf',
        )

        if Document_process._finalize_with_uno(input_file, pdf_output):
            return

        try:
            libre_path = Document_process._find_libreoffice()
        except FileNotFoundError as e:
            logger.error(str(e))
            return

        commandStrings = [
            libre_path, "--headless", "--convert-to",
            "pdf", "--outdir", output_folder, input_file
        ]
        try:
            retCode = subprocess.call(commandStrings, timeout=120)
            if retCode == 0:
                logger.info('Conversión a PDF exitosa: %s', input_file)
            else:
                logger.error('Error en conversión a PDF. Código: %d', retCode)
        except subprocess.TimeoutExpired:
            logger.error('Timeout al convertir a PDF: %s', input_file)
        except Exception as e:
            logger.error('Error inesperado en conversión: %s', e)

    # ── Paragraph helpers ──────────────────────────────────────────────

    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    @staticmethod
    def llenar_campos(replacements, document):
        """Reemplaza placeholders preservando el formato de los runs.

        No basta con buscar la clave dentro de cada run por separado: Word
        parte el texto de un párrafo en varios <w:r> por motivos internos
        (ediciones previas, distinto rsid, etc.), y algunas plantillas traen
        placeholders como '[teacher]' repartidos en dos runs ('[' por un
        lado, 'teacher]' por otro). En ese caso ningún run contiene la clave
        completa y el placeholder se queda tal cual, sin reemplazar. Por eso
        aquí se arma el texto completo del párrafo y se ubica la clave ahí,
        sin importar en cuántos runs esté repartida.
        """
        for paragraph in document.paragraphs:
            for key, value in replacements.items():
                Document_process._replace_in_paragraph(paragraph, key, value)

    @staticmethod
    def _replace_in_paragraph(paragraph, key, value):
        """Sustituye todas las ocurrencias de `key` en un párrafo, aunque
        estén repartidas entre varios runs. El formato que se conserva es
        el del run donde empieza la coincidencia."""
        while True:
            runs = paragraph.runs
            if not runs:
                return

            full_text = ''.join(r.text for r in runs)
            pos = full_text.find(key)
            if pos == -1:
                return
            end = pos + len(key)

            cursor = 0
            start_run_idx = start_offset = None
            end_run_idx = end_offset = None
            for i, r in enumerate(runs):
                r_start, r_end = cursor, cursor + len(r.text)
                if start_run_idx is None and r_start <= pos < r_end:
                    start_run_idx, start_offset = i, pos - r_start
                if r_start < end <= r_end:
                    end_run_idx, end_offset = i, end - r_start
                    break
                cursor = r_end
            if start_run_idx is None or end_run_idx is None:
                return

            if start_run_idx == end_run_idx:
                r = runs[start_run_idx]
                r.text = r.text[:start_offset] + value + r.text[end_offset:]
            else:
                first, last = runs[start_run_idx], runs[end_run_idx]
                first.text = first.text[:start_offset] + value
                for i in range(start_run_idx + 1, end_run_idx):
                    runs[i].text = ''
                last.text = last.text[end_offset:]

    @staticmethod
    def capitalizar_frases(cadena):
        palabras = cadena.split()
        palabras_capitalizadas = []
        skip_words = ['de', 'del', 'la', 'las', 'los', 'y', 'para']
        for palabra in palabras:
            if palabra.lower() in skip_words and palabras_capitalizadas:
                palabras_capitalizadas.append(palabra.lower())
            else:
                palabras_capitalizadas.append(palabra.capitalize())
        return ' '.join(palabras_capitalizadas)

    @staticmethod
    def generate_random_code(length=2):
        letters_and_digits = string.ascii_letters + string.digits
        return ''.join(random.choice(letters_and_digits) for i in range(length))

    @staticmethod
    def docx_replace(doc, old_text, new_text):
        for p in doc.paragraphs:
            if old_text in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        text = inline[i].text.replace(old_text, new_text)
                        inline[i].text = text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    Document_process.docx_replace(cell, old_text, new_text)

    @staticmethod
    def remove_file(file_path):
        base_path = os.path.splitext(file_path)[0]
        for extension in ['.pdf', '.docx']:
            try:
                os.remove(base_path + extension)
                logger.info('Archivo eliminado: %s%s', base_path, extension)
            except Exception as error:
                logger.warning("Error eliminando archivo: %s", error)

    # ── Posicionamiento absoluto en la portada ─────────────────────────

    @staticmethod
    def _find_paragraph(document, marker, exact=False):
        """Ubica un párrafo por su marcador de placeholder ('[title]', etc.)
        en vez de por índice fijo. Un índice numérico (document.paragraphs[17])
        se desalinea apenas cambia la estructura de arriba: por ejemplo,
        insert_logo() antepone un párrafo nuevo y corre todos los índices
        posteriores en +1. Buscar por el texto del marcador es inmune a eso."""
        for p in document.paragraphs:
            if (p.text == marker) if exact else (marker in p.text):
                return p
        return None

    @staticmethod
    def _anchor_paragraph_to_page(document, paragraph, y_align, x_align='center'):
        """Fija el párrafo en una posición absoluta de la página (marco
        clásico de Word/LibreOffice: w:framePr) y lo saca del flujo normal.

        Se usa para el título de portada (siempre centrado verticalmente) y
        la línea de fecha/lugar (siempre al pie), que antes se lograban
        contando párrafos en blanco antes/después. Ese conteo fijo se
        descuadraba con cualquier variación de contenido: la presencia o no
        del logo, el nombre de la universidad ocupando una o dos líneas, o
        cuántos integrantes se llenaron. Con el párrafo anclado a la página,
        su posición ya no depende de nada de lo anterior.
        """
        section = document.sections[0]
        width = section.page_width - section.left_margin - section.right_margin

        pPr = paragraph._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:framePr')):
            pPr.remove(old)

        frame = OxmlElement('w:framePr')
        frame.set(qn('w:w'), str(width))
        frame.set(qn('w:h'), '720')
        frame.set(qn('w:hRule'), 'auto')
        frame.set(qn('w:hAnchor'), 'page')
        frame.set(qn('w:vAnchor'), 'page')
        frame.set(qn('w:xAlign'), x_align)
        frame.set(qn('w:yAlign'), y_align)
        # w:framePr debe ir antes que w:tabs/w:spacing/w:jc dentro de pPr
        # (orden que exige el esquema CT_PPrBase); estos párrafos no traen
        # pStyle/keepNext/pageBreakBefore, así que va siempre primero.
        pPr.insert(0, frame)

    # ── TOC (Tabla de Contenido) ───────────────────────────────────────

    @staticmethod
    def _has_drawing(paragraph):
        """True si el párrafo trae una imagen/dibujo (aunque no tenga texto)."""
        return bool(
            paragraph._p.findall('.//' + qn('w:drawing'))
            or paragraph._p.findall('.//' + qn('w:pict'))
        )

    # Constantes calibradas empíricamente contra la plantilla real (tamaño
    # Carta, márgenes de 1"): posición Y (en puntos, desde el tope de la
    # página) donde termina el bloque de encabezado según haya o no logo y
    # según el nombre de la universidad quepa en una o dos líneas.
    _HEADER_END_BASE_PT = 150       # sin logo, nombre en 1 línea
    _HEADER_END_WRAP_EXTRA_PT = 16  # nombre en 2 líneas
    _HEADER_END_LOGO_EXTRA_PT = 86  # logo presente
    _HEADER_NAME_WRAP_THRESHOLD = 50  # a partir de este largo, se asume 2 líneas
    # Y mínima a la que puede empezar el pie sin invadir el título (anclado
    # al centro exacto de la página, ~396pt), con margen para títulos de
    # hasta 2-3 líneas.
    _TITLE_CLEAR_ZONE_PT = 440
    _MIN_SPACER_PT = 20

    @staticmethod
    def _trim_cover_spacers(document, title_para, has_logo=False, university_name=''):
        """Reemplaza los 22 párrafos en blanco que rodean el título por UNO
        solo, con la altura exacta que hace falta para que el pie de página
        (docente/integrantes) nunca choque con el título ni se desborde a
        una segunda página.

        Esos 22 párrafos eran el mecanismo original para 'empujar' el
        título al centro y el pie hacia abajo contando líneas a mano. Con el
        título y la fecha ya anclados a una posición absoluta de la página
        (ver _anchor_paragraph_to_page), un conteo FIJO de líneas en blanco
        no puede servir a la vez a los dos casos límite: si es lo bastante
        grande para no chocar con el título cuando el encabezado es corto
        (sin logo, nombre corto), sobra espacio y empuja el pie fuera de la
        página cuando el encabezado además es alto (con logo, nombre largo)
        + hay muchos integrantes. Por eso se calcula la altura del espaciador
        en función de lo que sí se conoce en el momento de generar el
        documento (si hay logo, si el nombre es largo), de modo que el pie
        arranque siempre en el mismo punto seguro sin importar el encabezado.
        """
        all_paragraphs = document.paragraphs
        title_idx = next(
            i for i, p in enumerate(all_paragraphs) if p._p is title_para._p
        )

        def is_removable_blank(p):
            return p.text.strip() == '' and not Document_process._has_drawing(p)

        before = [p for p in all_paragraphs[:title_idx] if is_removable_blank(p)]
        after = [p for p in all_paragraphs[title_idx + 1:] if is_removable_blank(p)]
        blanks = before + after
        if not blanks:
            return

        spacer = blanks[0]
        for p in blanks[1:]:
            Document_process.delete_paragraph(p)

        header_end = Document_process._HEADER_END_BASE_PT
        if len(university_name) >= Document_process._HEADER_NAME_WRAP_THRESHOLD:
            header_end += Document_process._HEADER_END_WRAP_EXTRA_PT
        if has_logo:
            header_end += Document_process._HEADER_END_LOGO_EXTRA_PT

        needed_pt = max(
            Document_process._MIN_SPACER_PT,
            Document_process._TITLE_CLEAR_ZONE_PT - header_end,
        )
        spacer.paragraph_format.space_after = Pt(needed_pt)

    @staticmethod
    def set_outline_level(paragraph, level):
        """Marca el nivel de esquema del párrafo (0 = nivel 1, 1 = nivel 2).

        Esto es lo que realmente lee el campo TOC para armar el índice. Se
        aplica DIRECTO sobre el párrafo y no sólo mediante el estilo, porque
        las plantillas de este proyecto no traen estilos 'Heading' y los que
        python-docx crea con add_style() nacen como estilos personalizados
        sin nivel de esquema: el índice quedaba vacío ('no se encontraron
        entradas de tabla de contenido') aunque el texto se viera en negrita.
        """
        pPr = paragraph._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:outlineLvl')):
            pPr.remove(old)
        outline = OxmlElement('w:outlineLvl')
        outline.set(qn('w:val'), str(level))
        pPr.append(outline)

    @staticmethod
    def _configure_heading_styles(document):
        """Crea/ajusta los estilos Heading 1 y Heading 2 con nivel de esquema."""
        for heading_name, level, size in (('Heading 1', 0, 14), ('Heading 2', 1, 12)):
            try:
                style = document.styles[heading_name]
            except KeyError:
                # Si no existe, crearlo basándose en Normal
                from docx.enum.style import WD_STYLE_TYPE
                style = document.styles.add_style(heading_name, WD_STYLE_TYPE.PARAGRAPH)

            font = style.font
            font.name = 'Arial'
            font.size = Pt(size)
            font.bold = True
            font.color.rgb = RGBColor(0, 0, 0)

            # El estilo también debe declarar su nivel de esquema, si no Word
            # lo trata como un párrafo normal en negrita y no lo indexa.
            style_el = style.element
            pPr = style_el.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                style_el.append(pPr)
            for old in pPr.findall(qn('w:outlineLvl')):
                pPr.remove(old)
            outline = OxmlElement('w:outlineLvl')
            outline.set(qn('w:val'), str(level))
            pPr.append(outline)

    @staticmethod
    def add_toc_page(document):
        """Inserta una página de Tabla de Contenido usando campos OOXML nativos."""
        # Salto de página explícito ANTES del título. Sin esto, "Índice"
        # aparecía en la posición donde el flujo normal de texto terminara
        # de llenar la página de portada: si el contenido de la portada era
        # corto, quedaba pegado justo debajo de la fecha en la misma página
        # 1; si era largo, se corría más abajo en la página 2. Un salto de
        # página real garantiza que siempre empiece arriba de su propia
        # página, sin importar cuánto contenido variable tenga la portada.
        document.add_page_break()

        # Título "Índice"
        p_title = document.add_paragraph()
        p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_title.paragraph_format.line_spacing = Pt(21)
        run_title = p_title.add_run('Índice')
        run_title.bold = True
        run_title.font.size = Pt(16)
        run_title.font.name = 'Arial'

        # Espacio
        document.add_paragraph('')

        # Campo TOC nativo de Word/LibreOffice
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = Pt(21)

        # Begin field char. w:dirty le pide a Word que recalcule el campo al
        # abrir el documento, sin esperar a que el usuario pulse F9.
        run1 = paragraph.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        fldChar_begin.set(qn('w:dirty'), 'true')
        run1._r.append(fldChar_begin)

        # Instruction text
        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._r.append(instrText)

        # Separate field char
        run3 = paragraph.add_run()
        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fldChar_sep)

        # Placeholder text (shown before TOC is updated)
        run4 = paragraph.add_run('Actualice este campo para ver el índice (F9)')
        run4.font.color.rgb = RGBColor(128, 128, 128)
        run4.font.size = Pt(11)

        # End field char
        run5 = paragraph.add_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run5._r.append(fldChar_end)

        # Salto de página después del TOC
        document.add_page_break()

    @staticmethod
    def set_update_fields(document):
        """Fuerza la actualización automática de campos al abrir el documento."""
        settings_element = document.settings.element
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings_element.append(update_fields)

    # ── Logo universitario ─────────────────────────────────────────────

    @staticmethod
    def _normalize_logo_name(university_name):
        """Normaliza el nombre de la universidad a un nombre de archivo."""
        # Quitar acentos
        nfkd = unicodedata.normalize('NFKD', university_name)
        ascii_name = ''.join(c for c in nfkd if not unicodedata.combining(c))
        # Minúsculas y reemplazar espacios
        normalized = ascii_name.lower().strip()
        normalized = re.sub(r'[^a-z0-9]+', '_', normalized)
        normalized = normalized.strip('_')
        return normalized

    @staticmethod
    def insert_logo(document, university_name, logo_dir='static/logos'):
        """Inserta el logo de la universidad centrado en la parte superior del
        documento. Devuelve True si lo encontró e insertó, False si no
        (se usa para calcular cuánto espacio reservar en _trim_cover_spacers)."""
        if not university_name or not university_name.strip():
            return False

        normalized = Document_process._normalize_logo_name(university_name)
        logo_path = None

        # Buscar el logo con distintas extensiones
        for ext in ['png', 'jpg', 'jpeg', 'webp']:
            candidate = os.path.join(logo_dir, f'{normalized}.{ext}')
            if os.path.isfile(candidate):
                logo_path = candidate
                break

        if logo_path is None:
            logger.info('Logo no encontrado para universidad: %s (buscado como: %s)',
                        university_name, normalized)
            return False

        try:
            # Insertar el logo al principio del documento, centrado
            # Tomamos el primer párrafo y lo usamos como ancla
            first_para = document.paragraphs[0]

            # Crear un nuevo párrafo ANTES del primer párrafo
            new_para = OxmlElement('w:p')
            # Configurar alineación centrada
            pPr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
            new_para.append(pPr)

            # Insertar antes del primer párrafo
            first_para._element.addprevious(new_para)

            # Ahora acceder al párrafo insertado a través de python-docx
            from docx.text.paragraph import Paragraph
            inserted_para = Paragraph(new_para, document)
            run = inserted_para.add_run()
            run.add_picture(logo_path, width=Cm(3))

            logger.info('Logo insertado: %s', logo_path)
            return True
        except Exception as e:
            logger.error('Error insertando logo: %s', e)
            return False

    # ── Generación de párrafos con estilos ─────────────────────────────

    # Longitud máxima que puede tener una línea para considerarse subtítulo
    SUBTITLE_MAX_LEN = 110

    @staticmethod
    def _clean_markdown(text):
        """Quita los marcadores markdown que el modelo emite de vez en cuando."""
        text = re.sub(r'^\s*#{1,6}\s*', '', text.strip())
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'^\*+\s*|\s*\*+$', '', text)
        return text.replace('\r', ' ').replace('\n', ' ').strip()

    @staticmethod
    def _is_subtitle(raw_line):
        """Decide si una línea del texto generado es un subtítulo.

        El modelo devuelve los bloques separados por saltos de línea, sin
        marcar cuáles son subtítulos, así que se deducen: son líneas cortas,
        que no cierran como una oración normal (un párrafo casi siempre
        termina en punto) o que vienen marcadas en markdown.
        """
        text = raw_line.strip()
        if not text:
            return False

        # Marcado explícito del modelo: '## Título' o '**Título**'
        if re.match(r'^\s*#{1,6}\s+\S', text):
            return True
        if re.match(r'^\*\*[^*]+\*\*[:.]?$', text):
            return True

        text = Document_process._clean_markdown(text)
        if not text or len(text) > Document_process.SUBTITLE_MAX_LEN:
            return False
        if text.endswith(':'):
            return True

        # Un párrafo normal termina con puntuación de cierre. Los paréntesis y
        # comillas finales no cuentan: un subtítulo puede acabar en ')', como
        # 'Capa 1 (Layer 1): La Cadena Principal (On-Chain)', así que se miran
        # descartando esos cierres.
        core = text.rstrip(')]}"”\'’')
        if core and core[-1] in '.;,':
            return False

        # Varias oraciones seguidas => es un párrafo, no un subtítulo
        if re.search(r'\.\s+\S', text):
            return False
        return True

    @staticmethod
    def _split_blocks(body):
        """Separa el texto generado en bloques (párrafos y subtítulos).

        El prompt le pide al modelo separar cada párrafo con '\\n\\n\\n'. Antes
        se hacía body.split('\\n') descartando las líneas vacías, con lo que
        esa separación se perdía por completo y todo quedaba pegado.
        """
        # Se corta por cualquier racha de saltos: el modelo no es consistente
        # (a veces usa '\n\n\n', a veces uno solo) y en este formato nunca
        # parte un mismo párrafo en varias líneas.
        return [b.strip() for b in re.split(r'[\r\n]+', body) if b.strip()]

    @staticmethod
    def parrafos(body, document, topic, flag):
        """Agrega una sección al documento: título, subtítulos y párrafos.

        El título de sección va con nivel de esquema 1 y cada subtítulo
        detectado con nivel 2, que es lo que hace que ambos aparezcan en el
        índice.
        """
        if body == '':
            return

        # ── Título de la sección (nivel 1 del índice) ──
        p = document.add_paragraph(topic)
        try:
            p.style = document.styles['Heading 1']
        except KeyError:
            pass  # Si no existe Heading 1, el formato manual de abajo basta
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.line_spacing = Pt(21)
        p.paragraph_format.space_after = Pt(18)
        Document_process.set_outline_level(p, 0)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        # ── Cuerpo ──
        for block in Document_process._split_blocks(body):
            text = Document_process._clean_markdown(block)
            if not text:
                continue

            if Document_process._is_subtitle(block):
                sp = document.add_paragraph(text)
                try:
                    sp.style = document.styles['Heading 2']
                except KeyError:
                    pass
                sp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                sp.paragraph_format.line_spacing = Pt(21)
                sp.paragraph_format.space_before = Pt(12)
                sp.paragraph_format.space_after = Pt(6)
                Document_process.set_outline_level(sp, 1)
                for run in sp.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                bp = document.add_paragraph(text)
                bp.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                bp.paragraph_format.line_spacing = Pt(21)
                # Separación real entre párrafos (antes quedaban pegados)
                bp.paragraph_format.space_after = Pt(12)
                for run in bp.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

        if flag:
            document.add_page_break()

    # ── Underline helpers ──────────────────────────────────────────────

    @staticmethod
    def underline_words_in_first_page(doc, words):
        # Se recorren TODOS los párrafos existentes en este momento, no un
        # rango de índices fijo (24-37): en esta etapa del pipeline el
        # documento sólo tiene los párrafos de la portada (el índice y el
        # contenido se agregan después), así que no hace falta acotar nada.
        # Un rango fijo se desalineaba en cuanto insert_logo() anteponía un
        # párrafo (el logo) y corría todo lo demás un índice hacia abajo.
        #
        # Sólo se reconstruyen los runs de párrafos que SÍ contienen alguna
        # palabra clave. Es necesario filtrar así y no sólo iterar todos:
        # reconstruir runs implica llamar run.clear(), que borra cualquier
        # contenido del run, incluida una imagen si la lleva (como el run
        # del logo, que no tiene texto pero sí un <w:drawing>). Sin este
        # filtro, ese párrafo se vaciaba igual y el logo desaparecía.
        for paragraph in doc.paragraphs:
            if not any(word in paragraph.text for word in words):
                continue
            new_runs = []
            for run in paragraph.runs:
                found = False
                for word in words:
                    if word in run.text:
                        parts = run.text.split(word)
                        for part in parts[:-1]:
                            new_runs.append((part, run.bold, run.italic, run.underline))
                            new_runs.append((word, run.bold, run.italic, True))
                        new_runs.append((parts[-1], run.bold, run.italic, run.underline))
                        found = True
                        break
                if not found:
                    new_runs.append((run.text, run.bold, run.italic, run.underline))
            for run in paragraph.runs:
                run.clear()
            for text, bold, italic, underline in new_runs:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                run.underline = underline

    # ── Main orchestration ─────────────────────────────────────────────

    @staticmethod
    def fill_placeholders(docx_output, template_path, template_path2, replacements,
                           introduction, essay_content, conclusion, head_title, id,
                           university_name=''):
        if id == 'bach':
            words = ['DOCENTE:', 'ALUMNOS:', 'ALUMNO:', 'MATERIA:']
        else:
            words = ['DOCENTE:', 'ALUMNOS:', 'ALUMNO:', 'SECCION:',
                     'AÑO:', 'SEMESTRE:', 'TRIMESTRE:', 'MATERIA:']

        document = Document(template_path)

        # Insertar logo de la universidad (centrado arriba en portada)
        has_logo = Document_process.insert_logo(document, university_name)

        # Ubicar el título y la línea de fecha/lugar por su marcador, ANTES
        # de reemplazar los placeholders (mientras el texto sigue siendo
        # literalmente '[title]'/'[date]', que es inconfundible). Buscarlos
        # por índice fijo se rompía apenas insert_logo() anteponía un
        # párrafo y corría todo lo demás.
        title_para = Document_process._find_paragraph(document, '[title]', exact=True)
        date_para = Document_process._find_paragraph(document, '[date]')

        # Recortar los párrafos en blanco que ya no cumplen ningún propósito
        # de posicionamiento (ver _trim_cover_spacers) y reservar sólo el
        # espacio que hace falta según haya o no logo y el largo real del
        # nombre de la universidad. Debe ir antes de llenar_campos: usa el
        # texto vacío para detectar los blancos, y llenar_campos no los toca
        # de todas formas.
        if title_para is not None:
            Document_process._trim_cover_spacers(
                document, title_para,
                has_logo=has_logo, university_name=replacements.get('[u]', ''),
            )

        # Llenar campos de la plantilla
        Document_process.llenar_campos(replacements, document)
        Document_process.underline_words_in_first_page(document, words)

        # Configurar estilo Normal
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        # Configurar estilos de Heading para el TOC
        Document_process._configure_heading_styles(document)

        # Título principal: tamaño de fuente y anclado al centro exacto de
        # la página. Anclarlo lo saca del flujo normal, así que queda
        # centrado siempre, sin importar si hay logo arriba, si el nombre
        # de la universidad ocupa una o dos líneas, o cuántos integrantes
        # se listen debajo.
        if title_para is not None and title_para.runs:
            title_para.runs[0].font.size = Pt(17.5)
            Document_process._anchor_paragraph_to_page(document, title_para, y_align='center')
        else:
            logger.warning('No se pudo ubicar/anclar el título en la portada')

        # Línea de fecha y lugar: anclada al pie de la página, siempre en
        # la misma posición sin importar cuánto ocupe el bloque de
        # docente/integrantes que va justo encima.
        if date_para is not None:
            Document_process._anchor_paragraph_to_page(document, date_para, y_align='bottom')
        else:
            logger.warning('No se pudo ubicar/anclar la línea de fecha en la portada')

        has_content = essay_content != '' or introduction != '' or conclusion != ''

        if has_content:
            # Insertar página de Tabla de Contenido
            Document_process.add_toc_page(document)

        flagi = False
        if essay_content != '' or conclusion != '':
            flagi = True

        flage = False
        if conclusion != '':
            flage = True

        Document_process.parrafos(introduction, document, 'Introducción', flagi)
        Document_process.parrafos(essay_content, document, head_title, flage)
        Document_process.parrafos(conclusion, document, 'Conclusión', False)

        # Forzar actualización de campos (TOC) al abrir
        if has_content:
            Document_process.set_update_fields(document)

        head_title = head_title.replace(':', '_')
        document.save(docx_output)
        Document_process.convert(docx_output, 'output')