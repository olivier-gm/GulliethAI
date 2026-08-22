import os
import re
import shutil
import zipfile
import pytest
from docx import Document
from algorythms import Document_process

# Assume current working directory during tests is the project root
TEMPLATE_PATH = 'Templates/Universitario.docx'
# La plantilla real de producción (con los placeholders '[title]', '[date]',
# etc.), a diferencia de TEMPLATE_PATH que es un docx dummy de relleno.
REAL_TEMPLATE_PATH = 'input/plantilla.docx'
TEST_OUTPUT_DIR = 'tests/output'
TEST_OUTPUT_DOCX = os.path.join(TEST_OUTPUT_DIR, 'test_output.docx')
TEST_OUTPUT_PDF = os.path.join('output', 'test_output.pdf') # hardcoded in app logic

@pytest.fixture(autouse=True)
def setup_teardown():
    # Setup
    os.makedirs(TEST_OUTPUT_DIR, exist_ok=True)
    os.makedirs('output', exist_ok=True)
    
    # We must have a test template, if it doesn't exist we make a dummy one for tests
    if not os.path.exists(TEMPLATE_PATH):
        os.makedirs(os.path.dirname(TEMPLATE_PATH), exist_ok=True)
        doc = Document()
        # Add a dummy paragraph to avoid index out of bounds when replacing things on cover
        for i in range(40):
            doc.add_paragraph(f'Paragraph {i}')
        doc.save(TEMPLATE_PATH)
        
    yield
    
    # Teardown
    if os.path.exists(TEST_OUTPUT_DIR):
        shutil.rmtree(TEST_OUTPUT_DIR)
    
    for f in [TEST_OUTPUT_PDF]:
        if os.path.exists(f):
            os.remove(f)

def test_document_generation_with_toc_and_logo():
    replacements = {'[TITLE]': 'Test Title'}
    introduction = 'Test Intro\nSecond line'
    essay_content = 'Test Body'
    conclusion = 'Test Conclusion'
    
    # Run the main generation function with a university that has a test logo
    # Assuming 'Universidad Central de Venezuela' maps to a logo that exists or at least won't crash
    Document_process.fill_placeholders(
        docx_output=TEST_OUTPUT_DOCX, 
        template_path=TEMPLATE_PATH, 
        template_path2='', 
        replacements=replacements,
        introduction=introduction, 
        essay_content=essay_content, 
        conclusion=conclusion, 
        head_title='Test Heading', 
        id='uni',
        university_name='Universidad Central de Venezuela'
    )
    
    # Verify file was created
    assert os.path.exists(TEST_OUTPUT_DOCX), "Output DOCX was not generated."
    
    # Open the generated DOCX and inspect it
    doc = Document(TEST_OUTPUT_DOCX)
    
    # 1. Verify TOC field is present in the XML
    # Nota: doc.element.xml devuelve str, no bytes. Y si LibreOffice esta
    # instalado, el docx final pasa por el (para poblar el indice), que
    # reordena los switches del campo, asi que no se compara el literal exacto.
    xml_content = doc.element.xml
    assert 'w:fldChar' in xml_content, "TOC field was not added (fldChar missing)."

    instr = ' '.join(re.findall(r'<w:instrText[^>]*>([^<]*)</w:instrText>', xml_content))
    assert 'TOC' in instr, "TOC instruction text not found."
    assert '1-3' in instr, "El campo TOC deberia abarcar los niveles 1 a 3."

    # 2. Verify Headings are applied
    heading_found = False
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1' and 'Test Heading' in p.text:
            heading_found = True
            break
    assert heading_found, "Heading 1 style was not correctly applied to body topic."

    # 3. Los encabezados deben declarar nivel de esquema; sin esto el indice
    #    sale vacio ("no se encontraron entradas de tabla de contenido").
    #    Puede estar en el parrafo (document.xml) o, si LibreOffice reescribio
    #    el archivo, en la definicion del estilo (styles.xml).
    with zipfile.ZipFile(TEST_OUTPUT_DOCX) as z:
        styles_xml = z.read('word/styles.xml').decode('utf-8')
    assert 'w:outlineLvl' in xml_content or 'outlineLvl' in styles_xml, \
        "Los encabezados no declaran nivel de esquema: el indice saldria vacio."

def test_logo_insertion_unknown_university():
    # Should not crash if the university is unknown and has no logo
    Document_process.fill_placeholders(
        docx_output=TEST_OUTPUT_DOCX, 
        template_path=TEMPLATE_PATH, 
        template_path2='', 
        replacements={},
        introduction='', 
        essay_content='Body only', 
        conclusion='', 
        head_title='Title',
        id='uni',
        university_name='Unknown University XYZ'
    )
    assert os.path.exists(TEST_OUTPUT_DOCX), "Document should generate successfully even if logo is missing."


# ── Separacion de parrafos y deteccion de subtitulos ───────────────────
# El modelo separa los parrafos con '\n\n\n'; antes esa separacion se perdia
# y todo el texto quedaba pegado en el documento.

def test_split_blocks_respeta_la_separacion_del_modelo():
    body = "Primer parrafo.\n\n\nUn Subtitulo\n\n\nSegundo parrafo."
    assert Document_process._split_blocks(body) == [
        'Primer parrafo.', 'Un Subtitulo', 'Segundo parrafo.',
    ]


def test_split_blocks_con_saltos_simples():
    # Si el modelo no deja lineas en blanco, se separa por saltos simples
    body = "Primer parrafo.\nSegundo parrafo."
    assert Document_process._split_blocks(body) == [
        'Primer parrafo.', 'Segundo parrafo.',
    ]


def test_split_blocks_con_separadores_mezclados():
    # El modelo no es consistente: mezcla '\n' y '\n\n\n' en un mismo texto
    body = "Uno.\nDos.\n\n\nTres.\r\n\r\nCuatro."
    assert Document_process._split_blocks(body) == [
        'Uno.', 'Dos.', 'Tres.', 'Cuatro.',
    ]


@pytest.mark.parametrize('linea', [
    'Origen y Fundamentos del Bitcoin',
    'Adopcion Institucional y Marcos Regulatorios',
    'Membrana Plasmatica:',
    '**Subtitulo en negrita**',
    '## Subtitulo markdown',
    # Un subtitulo puede cerrar con parentesis: antes se descartaban y no
    # salian ni en negrita ni en el indice
    'Capa 1 (Layer 1): La Cadena Principal (On-Chain)',
    'Capa 2 (Layer 2): Soluciones de Escalabilidad (Off-Chain)',
])
def test_detecta_subtitulos(linea):
    assert Document_process._is_subtitle(linea)


@pytest.mark.parametrize('linea', [
    'El Bitcoin, introducido en 2008 bajo el seudonimo de Satoshi Nakamoto, representa la primera '
    'implementacion exitosa de una moneda digital descentralizada.',
    'La arquitectura combina criptografia asimetrica y teoria de juegos.',
    'Una frase corta que termina en punto.',
    # Cierra con parentesis pero la frase termina en punto: sigue siendo parrafo
    'Esto es un ejemplo entre parentesis (con cierre.)',
])
def test_no_confunde_parrafos_con_subtitulos(linea):
    assert not Document_process._is_subtitle(linea)


def test_limpia_marcadores_markdown():
    assert Document_process._clean_markdown('## Titulo') == 'Titulo'
    assert Document_process._clean_markdown('**Titulo**') == 'Titulo'


# ── Portada de una sola página: título centrado, fecha al pie ──────────
# El título y la fecha se anclan con w:framePr a una posición absoluta de
# la página (independiente del logo, el largo del nombre de la universidad
# o la cantidad de integrantes). Los 22 párrafos en blanco que antes
# simulaban ese centrado ya no hacen falta y se recortan a uno solo, cuya
# altura se calcula dinámicamente según las dimensiones de la página, la
# cantidad de líneas de detalle y la posición de la fecha.

def test_trim_cover_spacers_reduce_los_parrafos_en_blanco():
    doc = Document(REAL_TEMPLATE_PATH)
    title_para = Document_process._find_paragraph(doc, '[title]', exact=True)
    date_para = Document_process._find_paragraph(doc, '[date]')
    assert title_para is not None

    blancos_antes = sum(1 for p in doc.paragraphs if p.text.strip() == '')
    detail_lines = Document_process._count_detail_lines(doc, title_para, date_para)
    Document_process._trim_cover_spacers(
        doc, title_para, date_para=date_para,
        has_logo=False, university_name='UCV',
        detail_lines=detail_lines,
    )
    blancos_despues = sum(1 for p in doc.paragraphs if p.text.strip() == '')

    # Sólo debe quedar un párrafo en blanco (el espaciador calculado).
    assert blancos_despues == 1
    assert blancos_despues < blancos_antes


def test_trim_cover_spacers_mas_detalle_menos_spacer():
    """Cuando hay más líneas de detalle, el spacer se reduce para que
    todo quepa en una sola página."""
    doc1 = Document(REAL_TEMPLATE_PATH)
    title1 = Document_process._find_paragraph(doc1, '[title]', exact=True)
    date1 = Document_process._find_paragraph(doc1, '[date]')
    Document_process._trim_cover_spacers(
        doc1, title1, date_para=date1, detail_lines=3,
    )
    spacer1 = next(p for p in doc1.paragraphs if p.text.strip() == '')
    pt1 = spacer1.paragraph_format.space_after.pt

    doc2 = Document(REAL_TEMPLATE_PATH)
    title2 = Document_process._find_paragraph(doc2, '[title]', exact=True)
    date2 = Document_process._find_paragraph(doc2, '[date]')
    Document_process._trim_cover_spacers(
        doc2, title2, date_para=date2, detail_lines=10,
    )
    spacer2 = next(p for p in doc2.paragraphs if p.text.strip() == '')
    pt2 = spacer2.paragraph_format.space_after.pt

    # Con más líneas de detalle, el spacer debe ser menor
    assert pt2 < pt1


def test_trim_cover_spacers_spacer_siempre_positivo():
    """Incluso con muchas líneas de detalle, el spacer nunca baja de
    _MIN_SPACER_PT."""
    doc = Document(REAL_TEMPLATE_PATH)
    title = Document_process._find_paragraph(doc, '[title]', exact=True)
    date = Document_process._find_paragraph(doc, '[date]')
    Document_process._trim_cover_spacers(
        doc, title, date_para=date, detail_lines=50,
    )
    spacer = next(p for p in doc.paragraphs if p.text.strip() == '')
    assert spacer.paragraph_format.space_after.pt >= Document_process._MIN_SPACER_PT


def test_count_detail_lines():
    """Verifica que _count_detail_lines cuenta los párrafos con texto
    entre el título y la fecha."""
    doc = Document(REAL_TEMPLATE_PATH)
    title = Document_process._find_paragraph(doc, '[title]', exact=True)
    date = Document_process._find_paragraph(doc, '[date]')
    assert title is not None and date is not None
    count = Document_process._count_detail_lines(doc, title, date)
    # La plantilla tiene varias líneas de detalle (docente, alumno, etc.)
    assert count > 0


def test_anchor_paragraph_to_page_uses_y_pos():
    """Cuando se pasa y_pos, el frame debe usar w:y en vez de w:yAlign."""
    from docx.shared import Emu
    doc = Document(REAL_TEMPLATE_PATH)
    date_para = Document_process._find_paragraph(doc, '[date]')
    assert date_para is not None

    # Posición Y absoluta: 9_000_000 EMU ≈ 709pt (cerca del fondo de Carta)
    Document_process._anchor_paragraph_to_page(doc, date_para, y_pos=9_000_000)

    from lxml import etree
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    frame = date_para._p.find('.//w:framePr', ns)
    assert frame is not None
    assert frame.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}y') is not None
    assert frame.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}yAlign') is None


def test_insert_logo_devuelve_false_si_no_encuentra_el_logo():
    doc = Document(REAL_TEMPLATE_PATH)
    assert Document_process.insert_logo(doc, 'Universidad Inexistente XYZ') is False


def test_insert_logo_devuelve_true_si_lo_encuentra():
    doc = Document(REAL_TEMPLATE_PATH)
    assert Document_process.insert_logo(doc, 'Universidad Central de Venezuela') is True


def test_underline_words_no_borra_el_logo():
    """Regresión: reconstruir los runs de TODOS los párrafos (para
    subrayar palabras clave) borraba el run del logo aunque no tuviera
    texto, porque run.clear() elimina cualquier contenido, incluida una
    imagen. Ahora sólo se tocan párrafos que sí contienen alguna palabra."""
    doc = Document(REAL_TEMPLATE_PATH)
    Document_process.insert_logo(doc, 'Universidad Central de Venezuela')
    Document_process.underline_words_in_first_page(
        doc, ['DOCENTE:', 'ALUMNOS:', 'ALUMNO:', 'SECCION:', 'MATERIA:']
    )
    assert Document_process._has_drawing(doc.paragraphs[0])


def test_fill_placeholders_caso_extremo_no_lanza_excepcion():
    """Universidad de nombre largo, todos los campos al máximo permitido
    por el formulario y los 8 integrantes llenos: no debe romper la
    generación (aunque no podemos medir el conteo real de páginas sin
    LibreOffice instalado, al menos la construcción del documento no debe
    fallar)."""
    reps = {
        '[u]': 'U' * 80, '[area]': 'A' * 40, '[carrera]': 'C' * 40,
        '[city]': 'CARACAS, ', '[date]': '20 DE AGOSTO DE 2026',
        '[seccion]': 'SECCION: "X"', '[title]': 'TITULO DE PRUEBA',
        '[docente]': 'DOCENTE:', '[teacher]': 'T' * 60,
        '[asignatura]': 'M' * 60, '[asignatura2]': '', '[asignatura_t]': 'MATERIA: ',
        '[periodo]': '', '[academico]': 'SEMESTRE: ', '[academico2]': '', '[periodo2]': '',
        '[alumnos]': 'ALUMNOS:',
    }
    for i in range(1, 9):
        reps[f'[{i}]'] = 'ESTUDIANTE ' * 3
        reps[f'[id{i}]'] = ' C.I- 1234567890'

    out = os.path.join(TEST_OUTPUT_DIR, 'extreme.docx')
    Document_process.fill_placeholders(
        out, REAL_TEMPLATE_PATH, '', reps, '', '', '', 'Cuerpo', 'uni',
        university_name='Universidad Central de Venezuela',
    )
    assert os.path.exists(out)


def test_fill_placeholders_pocos_campos():
    """Con pocos campos (1 alumno, docente, materia), la portada no debe
    lanzar excepción y el spacer debe ser mayor que con muchos campos."""
    reps = {
        '[u]': 'UNERG', '[area]': '', '[carrera]': '',
        '[city]': 'SAN JUAN, ', '[date]': '20 DE AGOSTO DE 2026',
        '[seccion]': '', '[title]': 'MI TITULO',
        '[docente]': 'DOCENTE:', '[teacher]': 'JUAN GARCIA',
        '[asignatura]': 'MATEMATICA', '[asignatura2]': '', '[asignatura_t]': 'MATERIA: ',
        '[periodo]': '', '[academico]': '', '[academico2]': '', '[periodo2]': '',
        '[alumnos]': 'ALUMNO:',
    }
    reps['[1]'] = 'VANESSA MORO'
    reps['[id1]'] = ' C.I- 14981769'
    for i in range(2, 9):
        reps[f'[{i}]'] = ''
        reps[f'[id{i}]'] = ''

    out = os.path.join(TEST_OUTPUT_DIR, 'pocos_campos.docx')
    Document_process.fill_placeholders(
        out, REAL_TEMPLATE_PATH, '', reps, '', '', '', 'Titulo', 'uni',
        university_name='',
    )
    assert os.path.exists(out)
